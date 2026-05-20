"""Word document export: large screenshots + paginated transcript."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

OUTPUT_DIR = Path("outputs")

# US Letter with tight margins — maximize content for AI readability
PAGE_MARGIN = Inches(0.45)
PAGE_IMAGE_MAX_WIDTH = Inches(7.6)
PAGE_IMAGE_MAX_HEIGHT = Inches(9.2)
TIMESTAMP_FONT_SIZE = Pt(9)
TRANSCRIPT_HEADING_SIZE = Pt(10)
TRANSCRIPT_FONT_SIZE = Pt(8)
TRANSCRIPT_LINES_PER_PAGE = 56
TRANSCRIPT_CHARS_PER_LINE = 105


def _safe_stem(filename: str) -> str:
    stem = Path(filename).stem or "video"
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._")
    return stem or "video"


def default_docx_name(video_filename: str) -> str:
    return f"{_safe_stem(video_filename)}_{date.today().isoformat()}.docx"


def _set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN


def _image_display_width(image_path: Path) -> Inches:
    with Image.open(image_path) as img:
        width_px, height_px = img.size
    if width_px <= 0 or height_px <= 0:
        return PAGE_IMAGE_MAX_WIDTH

    max_w = PAGE_IMAGE_MAX_WIDTH.inches
    max_h = PAGE_IMAGE_MAX_HEIGHT.inches
    aspect = width_px / height_px
    display_w = max_w
    display_h = display_w / aspect
    if display_h > max_h:
        display_w = max_h * aspect
    return Inches(display_w)


def _estimate_line_units(text: str) -> int:
    if not text:
        return 0
    return max(1, (len(text) + TRANSCRIPT_CHARS_PER_LINE - 1) // TRANSCRIPT_CHARS_PER_LINE)


def _add_transcript_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.font.size = TRANSCRIPT_FONT_SIZE
    run.font.color.rgb = RGBColor(30, 30, 30)


def _add_paginated_transcript(doc: Document, transcript_text: str) -> None:
    text = transcript_text.strip()
    if not text:
        return

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return

    doc.add_page_break()
    heading = doc.add_heading("Transcript", level=1)
    if heading.runs:
        heading.runs[0].font.size = TRANSCRIPT_HEADING_SIZE

    used_on_page = 0

    for line in lines:
        units = _estimate_line_units(line)
        if used_on_page > 0 and used_on_page + units > TRANSCRIPT_LINES_PER_PAGE:
            doc.add_page_break()
            used_on_page = 0

        _add_transcript_paragraph(doc, line)
        used_on_page += units


def create_screenshots_docx(
    screenshots: list[dict],
    *,
    video_filename: str,
    output_path: str | Path | None = None,
    transcript_text: str = "",
) -> Path:
    """Create a Word document with full-page screenshots and a paginated transcript appendix."""
    if not screenshots:
        raise ValueError("No screenshots were generated, so a document cannot be created.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = Path(output_path) if output_path else OUTPUT_DIR / default_docx_name(video_filename)

    doc = Document()
    _set_page_margins(doc)
    rendered = 0

    for index, shot in enumerate(screenshots, start=1):
        image_path = Path(str(shot.get("path", "")))
        if not image_path.is_file():
            continue

        if rendered > 0:
            doc.add_page_break()

        doc.add_picture(str(image_path), width=_image_display_width(image_path))

        label = str(shot.get("label", "00:00:00"))
        ts = doc.add_paragraph()
        ts.paragraph_format.space_before = Pt(2)
        ts.paragraph_format.space_after = Pt(0)
        run = ts.add_run(label)
        run.font.size = TIMESTAMP_FONT_SIZE
        run.font.color.rgb = RGBColor(60, 60, 60)
        run.bold = True
        rendered += 1

    if rendered == 0:
        raise ValueError("None of the generated screenshot files could be read.")

    _add_paginated_transcript(doc, transcript_text)
    doc.save(out)
    return out


create_screenshots_pdf = create_screenshots_docx
_default_pdf_name = default_docx_name
